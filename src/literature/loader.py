from __future__ import annotations

import hashlib
import logging
from dataclasses import dataclass
from pathlib import Path


LOGGER = logging.getLogger(__name__)
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class PageText:
    page_number: int
    text: str


@dataclass
class ArticleDocument:
    doc_id: str
    path: str
    filename: str
    extension: str
    text: str
    pages: list[PageText] | None = None


def load_articles(articles_dir: str) -> list[ArticleDocument]:
    root = Path(articles_dir)
    if not root.exists():
        LOGGER.warning("Articles directory does not exist: %s", root)
        return []

    documents: list[ArticleDocument] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            document = _load_article(path)
        except Exception as exc:  # pragma: no cover - defensive around damaged files
            LOGGER.warning("Failed to read article %s: %s", path, exc)
            continue
        if not document.text.strip():
            LOGGER.warning("Skipping empty article: %s", path)
            continue
        documents.append(document)
    return documents


def _load_article(path: Path) -> ArticleDocument:
    extension = path.suffix.lower()
    pages: list[PageText] | None = None
    if extension == ".pdf":
        pages = _read_pdf_pages(path)
        text = "\n\n".join(page.text for page in pages)
    elif extension == ".docx":
        text = _read_docx(path)
    else:
        text = _read_text(path)

    return ArticleDocument(
        doc_id=_doc_id(path),
        path=str(path),
        filename=path.name,
        extension=extension,
        text=text,
        pages=pages,
    )


def _read_pdf_pages(path: Path) -> list[PageText]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is required to read PDF articles") from exc

    reader = PdfReader(str(path))
    pages: list[PageText] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # pragma: no cover - page-level PDF failures vary
            LOGGER.warning("Failed to extract page %s from %s: %s", index, path, exc)
            text = ""
        if text.strip():
            pages.append(PageText(page_number=index, text=text))
    return pages


def _read_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("python-docx is required to read DOCX articles") from exc

    document = docx.Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())


def _read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8-sig", errors="replace")


def _doc_id(path: Path) -> str:
    digest = hashlib.sha1(str(path.resolve()).encode("utf-8")).hexdigest()[:12]
    return f"article_{digest}"
